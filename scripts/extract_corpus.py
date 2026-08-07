"""Extract corpus items from documentation/archive/*.docx into draft YAMLs.

Phase B step 0 (spec §3.3, plan §Phase B).

Reads:  documentation/archive/voice_ai_test_corpus.docx  (primary source)
Writes: corpus/_draft_conversational.yaml
        corpus/_draft_narration.yaml
        corpus/_draft_unassigned.yaml    (items from cut use cases)

Output is prefixed with `_draft_` because it is NEVER the final corpus —
it is a starting point for manual curation (spec §3.3).

Section detection is by TEXT PATTERN, not docx style — the source docx
has no heading styles applied (every paragraph is `(no style)`).

Patterns:
  H1 use-case:  `^\\d+\\.\\s+<short title>`   (< 60 chars, avoids matching
                                                 numbered legal clauses)
  H2 stratum:   text starts with a stratum name (STRATUM_MAP)
  Item ID:      `^[A-Z]\\d+$`                  (`L01`, `S03`, ...)
                Paragraphs BETWEEN two item IDs form one multi-paragraph
                item (long stratum ships items that span several paragraphs).

`Use for:` lines and the top-of-doc metadata are dropped.

Usage:
    python scripts/extract_corpus.py
    python scripts/extract_corpus.py --docx path/to/other.docx

Requires: python-docx  (`uv add python-docx --optional dev` if missing)
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore[import-not-found]
except ImportError:
    print(
        "ERROR: python-docx not installed. Run:\n    uv add python-docx --optional dev",
        file=sys.stderr,
    )
    sys.exit(1)

import yaml

REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = REPO_ROOT / "documentation" / "archive" / "voice_ai_test_corpus.docx"
OUT_DIR = REPO_ROOT / "corpus"

# --- Use case routing ------------------------------------------------------
# Only two use cases are in scope per spec §3.1.
USE_CASE_MAP: dict[str, str] = {
    "conversational agent": "conversational",
    "customer support": "conversational",
    "book and long-form narration": "narration",
    "long-form narration": "narration",
}

# --- Stratum routing -------------------------------------------------------
STRATUM_MAP: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"^Short utterances?", re.I), "short"),
    (re.compile(r"^Medium paragraphs?", re.I), "medium"),
    (re.compile(r"^Long passages?", re.I), "long"),
    (re.compile(r"^Jargon battery|^Jargon terms?", re.I), "jargon"),
    (re.compile(r"^Edge cases?|^Edge battery", re.I), "edge"),
]

# H1 detection: numbered heading with a SHORT title (avoids matching
# numbered legal clauses like "1. GRANT OF LICENCE...").
H1_RE = re.compile(r"^\d+\.\s+(?P<title>.{3,60}?)$")

# Item ID markers (e.g. `L01`, `S03`, `J12`, `E05`, `M04`).
ID_RE = re.compile(r"^[A-Z]\d+$")

# --- Drop patterns ---------------------------------------------------------
DROP_PATTERNS: list[re.Pattern[str]] = [
    re.compile(r"^Use for:", re.I),
    re.compile(r"^\s*\(\d+\s*items?", re.I),
    re.compile(r"^Voice AI Evaluation$", re.I),
    re.compile(r"^Test Corpus", re.I),
    re.compile(r"^\d+\s*use cases?\s*\S+\s*\d+\s*short", re.I),  # roster tagline
]


def _classify_stratum_by_length(text: str) -> str:
    """Fallback when stratum can't be read from surrounding H2."""
    words = len(text.split())
    if words < 15:
        return "short"
    if words <= 60:
        return "medium"
    return "long"


def _use_case_for_h1(title: str) -> str | None:
    if not title:
        return None
    low = title.lower()
    for key, uc in USE_CASE_MAP.items():
        if key in low:
            return uc
    return None


def _stratum_for_h2(heading: str) -> str | None:
    for pat, s in STRATUM_MAP:
        if pat.search(heading):
            return s
    return None


def _should_drop(text: str) -> bool:
    return any(p.search(text) for p in DROP_PATTERNS)


@dataclass
class _RawItem:
    """One corpus item as extracted from the docx (may span multiple paragraphs)."""

    id_marker: str                # e.g. `L01`
    paragraphs: list[str] = field(default_factory=list)
    h1: str = ""                  # use-case section
    h2: str = ""                  # stratum section

    @property
    def text(self) -> str:
        return " ".join(self.paragraphs).strip()


@dataclass
class _Extracted:
    text: str
    use_case: str            # `conversational`, `narration`, or `_unassigned`
    stratum: str
    h1: str
    h2: str
    id_marker: str = ""


def _iter_items(docx_path: Path) -> list[_Extracted]:
    """Walk paragraphs, group multi-paragraph items by ID marker."""
    doc = Document(str(docx_path))
    h1_title = ""
    h2 = ""
    current_item: _RawItem | None = None
    raw_items: list[_RawItem] = []
    orphan_items: list[_Extracted] = []   # paragraphs outside an ID group

    def flush_current() -> None:
        nonlocal current_item
        if current_item and current_item.paragraphs:
            raw_items.append(current_item)
        current_item = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # H1 (use-case section)
        m = H1_RE.match(text)
        if m:
            flush_current()
            h1_title = m.group("title").strip()
            h2 = ""
            continue

        # H2 (stratum)
        if _stratum_for_h2(text) is not None:
            flush_current()
            h2 = text
            continue

        # Item ID marker
        if ID_RE.match(text):
            flush_current()
            current_item = _RawItem(id_marker=text, h1=h1_title, h2=h2)
            continue

        # Drop metadata lines
        if _should_drop(text):
            continue

        # Strip list-item prefixes ("1. ", "1) ", "- ", "* ", "• ")
        text = re.sub(r"^\s*\d+[.)]\s+", "", text)
        text = re.sub(r"^\s*[-•*]\s+", "", text)

        if current_item is not None:
            # Inside an item's ID group — this paragraph belongs to the current item
            current_item.paragraphs.append(text)
        else:
            # Orphan: content paragraph not preceded by an ID marker.
            # Treat as one item per paragraph, best-effort stratum guess.
            uc = _use_case_for_h1(h1_title) or "_unassigned"
            stratum = _stratum_for_h2(h2) or _classify_stratum_by_length(text)
            orphan_items.append(_Extracted(text=text, use_case=uc, stratum=stratum,
                                           h1=h1_title, h2=h2, id_marker=""))

    flush_current()

    # Consolidate raw items (multi-paragraph) into _Extracted
    extracted: list[_Extracted] = list(orphan_items)
    for ri in raw_items:
        uc = _use_case_for_h1(ri.h1) or "_unassigned"
        stratum = _stratum_for_h2(ri.h2) or _classify_stratum_by_length(ri.text)
        extracted.append(_Extracted(text=ri.text, use_case=uc, stratum=stratum,
                                     h1=ri.h1, h2=ri.h2, id_marker=ri.id_marker))

    return extracted


def _to_yaml_items(items: list[_Extracted], use_case: str) -> list[dict[str, Any]]:
    """Assign stable ids per stratum and serialize to CorpusItem shape."""
    stratum_counts: dict[str, int] = {}
    out: list[dict[str, Any]] = []
    prefix_map = {"short": "S", "medium": "M", "long": "L", "jargon": "J", "edge": "E", "probe": "P"}

    for item in items:
        if item.use_case != use_case:
            continue
        stratum_counts[item.stratum] = stratum_counts.get(item.stratum, 0) + 1
        prefix = prefix_map.get(item.stratum, "X")
        item_id = f"{prefix}{stratum_counts[item.stratum]:02d}"
        tags: list[str] = []
        if item.h1:
            tags.append(f"src:h1:{item.h1}")
        if item.h2:
            tags.append(f"src:h2:{item.h2}")
        if item.id_marker:
            tags.append(f"src:id:{item.id_marker}")
        out.append(
            {
                "id": item_id,
                "stratum": item.stratum,
                "text": item.text,
                "word_count": len(item.text.split()),
                "tags": tags,
            }
        )
    return out


def _to_unassigned_items(items: list[_Extracted]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for item in items:
        if item.use_case != "_unassigned":
            continue
        out.append(
            {
                "text": item.text,
                "word_count": len(item.text.split()),
                "stratum_guess": item.stratum,
                "h1": item.h1,
                "h2": item.h2,
                "id_marker": item.id_marker,
            }
        )
    return out


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract corpus items from a docx into draft YAMLs")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--out", type=Path, default=OUT_DIR)
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"ERROR: {args.docx} not found", file=sys.stderr)
        return 2

    args.out.mkdir(parents=True, exist_ok=True)
    extracted = _iter_items(args.docx)

    total = len(extracted)
    by_uc: dict[str, int] = {}
    for e in extracted:
        by_uc[e.use_case] = by_uc.get(e.use_case, 0) + 1

    print(f"Extracted {total} items from {args.docx.name}")
    print(f"  conversational: {by_uc.get('conversational', 0)}")
    print(f"  narration:      {by_uc.get('narration', 0)}")
    print(f"  unassigned:     {by_uc.get('_unassigned', 0)} (from cut use cases; reviewer triage)")

    for use_case in ("conversational", "narration"):
        payload = {
            "use_case": use_case,
            "items": _to_yaml_items(extracted, use_case),
        }
        out_path = args.out / f"_draft_{use_case}.yaml"
        with out_path.open("w", encoding="utf-8") as f:
            yaml.safe_dump(payload, f, sort_keys=False, allow_unicode=True, width=100)
        print(f"  wrote {out_path.relative_to(REPO_ROOT)}  ({len(payload['items'])} items)")

        strata_counts: dict[str, int] = {}
        for item in payload["items"]:
            strata_counts[item["stratum"]] = strata_counts.get(item["stratum"], 0) + 1
        for s in ("short", "medium", "long", "jargon", "edge"):
            print(f"      {s:8s} {strata_counts.get(s, 0)}")

    unassigned = _to_unassigned_items(extracted)
    if unassigned:
        out_path = args.out / "_draft_unassigned.yaml"
        with out_path.open("w", encoding="utf-8") as f:
            yaml.safe_dump({"unassigned": unassigned}, f, sort_keys=False, allow_unicode=True, width=100)
        print(f"  wrote {out_path.relative_to(REPO_ROOT)}  ({len(unassigned)} items)")

        from collections import Counter
        h1_counts = Counter(u["h1"] for u in unassigned)
        print("      by H1 section:")
        for h1, n in h1_counts.most_common():
            print(f"        {n:3d}  {h1 or '(no heading)'}")

    print(
        "\nDrafts written with `_draft_` prefix. Next steps (spec §3.3):\n"
        "  1. Manually curate/trim to 60 novel items per use case\n"
        "  2. Fix AI-generation artifacts; verify jargon items are actually hard\n"
        "  3. Confirm all names/amounts are synthetic (no real PII)\n"
        "  4. Add 15 famous public-domain probe items per use case\n"
        "  5. Save as `corpus/conversational.yaml` and `corpus/narration.yaml`\n"
        "  6. Select 10 items per use case for `corpus/variance_subset.yaml`\n"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
